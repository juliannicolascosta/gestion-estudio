from __future__ import annotations

import json
import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import time
import unicodedata
from datetime import date
from pathlib import Path
from typing import Callable

from .models import ADVANCED_FIELD_VARIABLES, AppSettings, Case, CompilationResult


APP_DIR = Path(os.getenv("APPDATA", Path.home())) / "GestorDocumental"
OFFICE_CACHE_DIR = Path(os.getenv("LOCALAPPDATA", APP_DIR)) / "GestorDocumental" / "conversion-cache"
CONFIG_NAME = "config.json"
CASE_METADATA = ".gestor-caso.json"
STUDY_LIBRARY_NAME = "00 - ACCESO RÁPIDO"
OFFICE_EXTENSIONS = {".doc", ".docx", ".odt", ".rtf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}
PDF_EXTENSIONS = OFFICE_EXTENSIONS | IMAGE_EXTENSIONS | {".pdf"}


class CompilationCancelled(RuntimeError):
    """Raised when the user asks to stop a compilation in progress."""


def _check_cancelled(cancelled: Callable[[], bool]) -> None:
    if cancelled():
        raise CompilationCancelled("Compilación cancelada.")


def safe_name(value: str) -> str:
    bad = '<>:"/\\|?*'
    cleaned = "".join("-" if char in bad else char for char in value)
    return " ".join(cleaned.split()).strip(" .-")


def repair_text(value: str) -> str:
    """Repair common UTF-8-as-Latin-1 mojibake left by early MVP builds."""
    if not any(marker in value for marker in ("Ã", "Â", "â")):
        return value
    try:
        repaired = value.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return value
    return repaired if repaired.count("Ã") < value.count("Ã") else value


def normalize_filename(value: str, extension: str | None = None) -> str:
    """Normalize unsafe punctuation and whitespace while preserving readable accents."""
    candidate = Path(value).name
    current_extension = Path(candidate).suffix.lower()
    stem = safe_name(Path(candidate).stem if current_extension else candidate)
    stem = stem or "ARCHIVO"
    selected_extension = extension if extension is not None else current_extension
    if selected_extension and not selected_extension.startswith("."):
        selected_extension = "." + selected_extension
    return stem + (selected_extension or "").lower()


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    number = 2
    while True:
        candidate = path.with_name(f"{path.stem} ({number}){path.suffix}")
        if not candidate.exists():
            return candidate
        number += 1


def versioned_path(path: Path) -> Path:
    """Return a visible presentation version instead of an ambiguous “(2)”."""
    if not path.exists():
        return path
    number = 2
    while True:
        candidate = path.with_name(f"{path.stem}_V{number}{path.suffix}")
        if not candidate.exists():
            return candidate
        number += 1


def template_variable_name(label: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(label))
    ascii_text = "".join(char for char in normalized if not unicodedata.combining(char))
    key = re.sub(r"[^A-Za-z0-9]+", "_", ascii_text).strip("_").upper()
    return key or "DATO"


def short_case_identifier(case: Case, metadata: dict[str, str] | None = None) -> str:
    metadata = metadata if metadata is not None else read_case_metadata(case)
    configured = metadata.get("Nombre corto para archivos", "").strip()
    if configured:
        return configured
    actor = metadata.get("Actor", "").strip()
    if not actor:
        return case.name.split(" c/ ", 1)[0].strip()
    primary = re.split(r"\s+(?:y|e)\s+(?:ot\.?|otros?)\b", actor, 1, flags=re.IGNORECASE)[0]
    if "," in primary:
        return primary.split(",", 1)[0].strip()
    company_markers = re.compile(
        r"\b(?:s\.?a\.?|s\.?r\.?l\.?|banco|sociedad|cooperativa|mutual|aseguradora)\b",
        re.IGNORECASE,
    )
    words = [word for word in primary.split() if word]
    if company_markers.search(primary):
        return " ".join(words[:4])
    return words[-1] if words else primary


def filename_component(value: str, maximum: int = 60) -> str:
    text = unicodedata.normalize("NFC", safe_name(value)).upper()
    text = re.sub(r"[^\w]+", "_", text, flags=re.UNICODE).strip("_")
    text = re.sub(r"_+", "_", text)
    return (text[:maximum].rstrip("_") or "SIN_DATO")


def writing_title_from_path(path: Path | None) -> str:
    if not path:
        return "PRESENTACION"
    stem = re.sub(r"^\d{4}-\d{2}-\d{2}\s*-\s*", "", path.stem).strip()
    return stem or "PRESENTACION"


def suggested_presentation_name(
    case: Case,
    writing: Path | None = None,
    value_date: date | None = None,
) -> str:
    metadata = read_case_metadata(case)
    identifier = filename_component(short_case_identifier(case, metadata), 40)
    title = filename_component(writing_title_from_path(writing), 60)
    stamp = (value_date or date.today()).isoformat()
    return f"{identifier}_{stamp}_{title}.pdf"


def ensure_bundled_writing_models(models_dir: Path) -> list[Path]:
    """Install new built-in models once while preserving every user edit."""
    bundled_dir = Path(__file__).with_name("bundled-models")
    if not bundled_dir.is_dir():
        return []
    installed = []
    for source in sorted(bundled_dir.glob("*.docx"), key=lambda path: path.name.casefold()):
        target = models_dir / source.name
        if not target.exists():
            shutil.copy2(source, target)
            installed.append(target)
    return installed


class SettingsStore:
    def __init__(self, app_dir: Path | None = None):
        self.app_dir = app_dir or APP_DIR
        self.config = self.app_dir / CONFIG_NAME
        self.models_dir = self.app_dir / "Modelos"
        self.base_template = self.models_dir / "Modelo base - Escrito nuevo.docx"
        self.app_dir.mkdir(parents=True, exist_ok=True)
        self.models_dir.mkdir(parents=True, exist_ok=True)
        ensure_bundled_writing_models(self.models_dir)
        self.settings = self.load()

    def load(self) -> AppSettings:
        try:
            payload = json.loads(self.config.read_text(encoding="utf-8"))
        except (FileNotFoundError, ValueError, TypeError):
            payload = {}
        professionals = [
            repair_text(str(item).strip())
            for item in payload.get("professionals", [])
            if str(item).strip()
        ]
        if not professionals:
            professionals = ["Profesional"]
        current = repair_text(str(payload.get("current_professional", "")).strip())
        if current not in professionals:
            current = professionals[0]
        roots_payload = payload.get("study_roots", [])
        roots = [
            Path(value)
            for value in roots_payload
            if isinstance(value, str) and value.strip()
        ]
        legacy_root = payload.get("study_root")
        if legacy_root and not roots:
            roots = [Path(legacy_root)]
        if not roots:
            legacy_cases = payload.get("cases", [])
            legacy_parents = {
                str(Path(row["path"]).parent)
                for row in legacy_cases
                if isinstance(row, dict) and row.get("path")
            }
            if len(legacy_parents) == 1:
                inferred = Path(legacy_parents.pop())
                # Never reinterpret the application source directory as the Study.
                if not ((inferred / "run.py").is_file() and (inferred / "gestor_documental").is_dir()):
                    roots = [inferred]
        deduplicated = []
        seen_roots = set()
        for root in roots:
            key = os.path.normcase(os.path.abspath(str(root)))
            if key not in seen_roots:
                seen_roots.add(key)
                deduplicated.append(root)
        roots = deduplicated
        active_value = payload.get("active_study_root") or legacy_root
        active = Path(active_value) if active_value else None
        if active not in roots:
            active = roots[0] if roots else None
        signer = payload.get("signer_path")
        return AppSettings(
            study_roots=roots,
            active_study_root=active,
            professionals=professionals,
            current_professional=current,
            signer_path=Path(signer) if signer else None,
        )

    def save(self):
        payload = {
            "study_root": str(self.settings.study_root) if self.settings.study_root else None,
            "study_roots": [str(path) for path in self.settings.study_roots],
            "active_study_root": (
                str(self.settings.study_root) if self.settings.study_root else None
            ),
            "professionals": self.settings.professionals,
            "current_professional": self.settings.current_professional,
            "signer_path": str(self.settings.signer_path) if self.settings.signer_path else None,
        }
        self.config.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def set_study_root(self, path: Path):
        self.add_study_root(path)

    def add_study_root(self, path: Path):
        path = Path(path)
        path.mkdir(parents=True, exist_ok=True)
        if path not in self.settings.study_roots:
            self.settings.study_roots.append(path)
        self.settings.active_study_root = path
        self.save()

    def set_active_study_root(self, path: Path):
        path = Path(path)
        if path not in self.settings.study_roots:
            raise ValueError("La ubicación no pertenece al Estudio.")
        self.settings.active_study_root = path
        self.save()

    def remove_study_root(self, path: Path):
        path = Path(path)
        self.settings.study_roots = [root for root in self.settings.study_roots if root != path]
        if self.settings.active_study_root == path:
            self.settings.active_study_root = (
                self.settings.study_roots[0] if self.settings.study_roots else None
            )
        self.save()

    def add_professional(self, name: str):
        name = " ".join(name.split()).strip()
        if name and name not in self.settings.professionals:
            self.settings.professionals.append(name)
        if name:
            self.settings.current_professional = name
        self.save()

    def set_professional(self, name: str):
        if name in self.settings.professionals:
            self.settings.current_professional = name
            self.save()

    def set_signer(self, path: Path | None):
        self.settings.signer_path = path
        self.save()


def list_cases(study_root: Path | None) -> list[Case]:
    if not study_root or not study_root.is_dir():
        return []
    return [
        Case(path)
        for path in sorted(study_root.iterdir(), key=lambda item: item.name.casefold())
        if path.is_dir()
        and not path.name.startswith(".")
        and path.name.casefold() != STUDY_LIBRARY_NAME.casefold()
    ]


def study_library_path(study_root: Path | None, create: bool = False) -> Path | None:
    if not study_root:
        return None
    path = study_root / STUDY_LIBRARY_NAME
    if create:
        path.mkdir(parents=True, exist_ok=True)
    return path


def create_case(study_root: Path, name: str) -> Case:
    normalized = safe_name(name)
    if not normalized:
        raise ValueError("Escribí un nombre para el caso.")
    path = study_root / normalized
    if path.exists():
        raise FileExistsError(f"Ya existe un caso llamado “{normalized}”.")
    case = Case(path)
    case.ensure()
    return case


def rename_case(case: Case, new_name: str) -> Case:
    normalized = safe_name(new_name)
    if not normalized:
        raise ValueError("Escribí un nombre para el caso.")
    target = case.path.with_name(normalized)
    if target == case.path:
        return case
    if target.exists():
        raise FileExistsError(f"Ya existe un caso llamado “{normalized}”.")
    case.path.rename(target)
    return Case(target)


def move_to_recycle_bin(paths: list[Path]) -> None:
    """Move files to the Windows Recycle Bin instead of deleting them permanently."""
    existing = [Path(path).resolve() for path in paths if Path(path).exists()]
    if not existing:
        return
    if os.name != "nt":
        raise OSError("La Papelera solo está disponible en la versión para Windows.")

    import ctypes
    from ctypes import wintypes

    class SHFILEOPSTRUCTW(ctypes.Structure):
        _fields_ = [
            ("hwnd", wintypes.HWND),
            ("wFunc", wintypes.UINT),
            ("pFrom", wintypes.LPCWSTR),
            ("pTo", wintypes.LPCWSTR),
            ("fFlags", ctypes.c_ushort),
            ("fAnyOperationsAborted", wintypes.BOOL),
            ("hNameMappings", ctypes.c_void_p),
            ("lpszProgressTitle", wintypes.LPCWSTR),
        ]

    sources = "\0".join(str(path) for path in existing) + "\0\0"
    operation = SHFILEOPSTRUCTW()
    operation.wFunc = 3  # FO_DELETE
    operation.pFrom = sources
    operation.fFlags = 0x0040 | 0x0010 | 0x0004 | 0x0400  # ALLOWUNDO, NOCONFIRMATION, SILENT, NOERRORUI
    result = ctypes.windll.shell32.SHFileOperationW(ctypes.byref(operation))
    if result or operation.fAnyOperationsAborted:
        raise OSError(f"Windows no pudo enviar los archivos a la Papelera (código {result}).")


def case_metadata_path(case: Case) -> Path:
    return case.path / CASE_METADATA


def read_case_metadata(case: Case) -> dict[str, str]:
    try:
        payload = json.loads(case_metadata_path(case).read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            return {}
        return {str(key): str(value) for key, value in payload.items()}
    except (FileNotFoundError, ValueError, TypeError):
        return {}


def save_case_metadata(case: Case, metadata: dict[str, str]):
    case.ensure()
    cleaned = {
        str(key): " ".join(str(value).split()).strip()
        for key, value in metadata.items()
        if str(value).strip()
    }
    metadata_path = case_metadata_path(case)
    temporary_path = metadata_path.with_suffix(metadata_path.suffix + ".tmp")
    if os.name == "nt" and metadata_path.exists():
        # Windows rejects opening a hidden file for replacement. Temporarily
        # clear the attribute, then restore it after the atomic write.
        subprocess.run(
            ["attrib", "-h", str(metadata_path)],
            check=False,
            capture_output=True,
        )
    try:
        temporary_path.write_text(
            json.dumps(cleaned, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        os.replace(temporary_path, metadata_path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()
        if os.name == "nt" and metadata_path.exists():
            subprocess.run(
                ["attrib", "+h", str(metadata_path)],
                check=False,
                capture_output=True,
            )


def _search_text(value: str) -> str:
    expanded = value.casefold().replace("nº", "numero").replace("n°", "numero")
    normalized = unicodedata.normalize("NFKD", expanded)
    return "".join(char for char in normalized if not unicodedata.combining(char))


def case_matches(case: Case, query: str) -> bool:
    tokens = [_search_text(token) for token in query.split() if token.strip()]
    if not tokens:
        return True
    metadata = read_case_metadata(case)
    haystack = _search_text(" ".join([case.name, *metadata.values()]))
    return all(token in haystack for token in tokens)


def open_file(path: Path):
    os.startfile(str(path))


def rename_case_file(source: Path, new_name: str) -> Path:
    extension = source.suffix
    target_name = normalize_filename(new_name, extension)
    target = source.with_name(target_name)
    if target == source:
        return source
    if target.exists():
        raise FileExistsError(f"Ya existe un archivo llamado “{target.name}”.")
    source.rename(target)
    return target


def rename_case_entry(source: Path, new_name: str) -> Path:
    if source.is_file():
        return rename_case_file(source, new_name)
    normalized = safe_name(new_name)
    if not normalized:
        raise ValueError("Escribí un nombre.")
    target = source.with_name(normalized)
    if target == source:
        return source
    if target.exists():
        raise FileExistsError(f"Ya existe una carpeta llamada “{target.name}”.")
    source.rename(target)
    return target


def can_convert_to_pdf(path: Path) -> bool:
    return path.suffix.lower() in OFFICE_EXTENSIONS | IMAGE_EXTENSIONS


def import_file(case: Case, source: Path, name: str, convert_to_pdf: bool = False) -> Path:
    if not source.is_file():
        raise FileNotFoundError(f"No encontramos {source.name}.")
    case.ensure()
    if convert_to_pdf:
        if not can_convert_to_pdf(source):
            raise ValueError(f"{source.name} no necesita conversión a PDF.")
        with tempfile.TemporaryDirectory(prefix="gestor-documental-") as directory:
            converted = to_pdf(source, Path(directory))
            target = unique_path(case.path / normalize_filename(name, ".pdf"))
            shutil.copy2(converted, target)
            return target
    extension = source.suffix.lower()
    target = unique_path(case.path / normalize_filename(name, extension))
    if source.resolve() == target.resolve():
        return source
    shutil.copy2(source, target)
    return target


def import_directory(case: Case, source: Path, name: str | None = None) -> Path:
    """Copy a user-selected directory as-is; the app never creates it implicitly."""
    if not source.is_dir():
        raise NotADirectoryError(f"No encontramos la carpeta {source.name}.")
    case.ensure()
    normalized = safe_name(name or source.name)
    if not normalized:
        raise ValueError("Escribí un nombre para la carpeta.")
    target = unique_path(case.path / normalized)
    shutil.copytree(source, target)
    return target


def _set_spanish_argentina(document):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for style_name in ("Normal", "Default Paragraph Font"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        run_properties = style.element.get_or_add_rPr()
        language = run_properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            run_properties.append(language)
        language.set(qn("w:val"), "es-AR")
        language.set(qn("w:eastAsia"), "es-AR")
        language.set(qn("w:bidi"), "es-AR")

    settings = document.settings.element
    theme_language = settings.find(qn("w:themeFontLang"))
    if theme_language is None:
        theme_language = OxmlElement("w:themeFontLang")
        settings.append(theme_language)
    theme_language.set(qn("w:val"), "es-AR")


def ensure_default_writing_template(path: Path) -> Path:
    from docx import Document

    bundled = Path(__file__).with_name("default-writing-template.docx")
    if path.is_file():
        try:
            existing = Document(path)
            legacy_text = [paragraph.text.strip() for paragraph in existing.paragraphs]
            generated_defaults = (
                ["TÍTULO DEL ESCRITO.", "Señor/a Juez/a:", "[Complete aquí el escrito]"],
                ["{{TITULO}}.", "{{CARATULA}}", "Señor/a Juez/a:", "[Complete aquí el escrito]"],
            )
            if legacy_text not in generated_defaults:
                return path
        except Exception:
            return path

    path.parent.mkdir(parents=True, exist_ok=True)
    if bundled.is_file():
        shutil.copy2(bundled, path)
        return path
    document = Document()
    _set_spanish_argentina(document)
    section = document.sections[0]
    section.top_margin = section.bottom_margin = 1134000
    paragraph = document.add_paragraph()
    paragraph.alignment = 2
    run = paragraph.add_run("{{TITULO}}.")
    run.bold = True
    caption = document.add_paragraph()
    caption.alignment = 2
    caption_run = caption.add_run("{{CARATULA}}")
    caption_run.bold = True
    document.add_paragraph("Señor/a Juez/a:")
    document.add_paragraph("[Complete aquí el escrito]")
    document.save(path)
    return path


SPANISH_MONTHS = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)


def spanish_long_date(value: date) -> str:
    return f"{value.day} de {SPANISH_MONTHS[value.month - 1]} de {value.year}"


def writing_template_values(
    case: Case,
    title: str,
    professional: str = "",
) -> dict[str, str]:
    metadata = read_case_metadata(case)
    actor = metadata.get("Actor", "").strip()
    defendant = metadata.get("Demandado", "").strip()
    cause = metadata.get("Causa", "").strip()
    case_number = (
        metadata.get("CUIJ", "").strip()
        or metadata.get("Número de expediente", "").strip()
    )
    caption = actor
    if defendant:
        caption = f"{caption} c/ {defendant}" if caption else defendant
    if cause:
        caption = f"{caption} s/ {cause}" if caption else cause
    caption = (caption or case.name).upper()
    metadata_lawyer = metadata.get("Abogado", "").strip()
    lawyer = metadata_lawyer or professional.strip()
    lawyer = re.sub(
        r"^\s*(?:dr|dra|doctor|doctora)\.?\s+",
        "",
        lawyer,
        flags=re.IGNORECASE,
    ).strip().upper()
    values = {
        "TITULO": safe_name(title).upper(),
        "CARATULA": caption,
        "CARÁTULA": caption,
        "CASO": case.name,
        "ACTOR": actor,
        "DEMANDADO": defendant,
        "CAUSA": cause,
        "CUIJ": case_number,
        "NUMERO_EXPEDIENTE": case_number,
        "NÚMERO_EXPEDIENTE": case_number,
        "EXPEDIENTE": case_number,
        "RADICACION": metadata.get("Radicación", "").strip(),
        "RADICACIÓN": metadata.get("Radicación", "").strip(),
        "ABOGADO": lawyer,
        "CONTRAPARTE": metadata.get("Contraparte", "").strip(),
        "PROFESIONAL": lawyer,
        "FECHA": date.today().strftime("%d/%m/%Y"),
        "FECHA_ISO": date.today().isoformat(),
        "FECHA_EXTENSA": spanish_long_date(date.today()),
    }
    for label, value in metadata.items():
        key = ADVANCED_FIELD_VARIABLES.get(label, template_variable_name(label))
        values.setdefault(key, str(value).strip())
    values["ACTOR_CORTO"] = short_case_identifier(case, metadata)
    values["NOMBRE_CORTO"] = values["ACTOR_CORTO"]
    values["CUIJ_COMPLETO"] = f" (CUIJ N° {values['CUIJ']})" if values["CUIJ"] else ""
    values["PROFESIONAL_MAYUSCULAS"] = lawyer
    return {f"{{{{{key}}}}}": value for key, value in values.items()}


def _replace_template_paragraph(paragraph, replacements: dict[str, str]) -> None:
    for placeholder, value in replacements.items():
        while paragraph.runs:
            combined = "".join(run.text for run in paragraph.runs)
            start = combined.find(placeholder)
            if start < 0:
                break
            end = start + len(placeholder)
            cursor = 0
            start_run = end_run = None
            start_offset = end_offset = 0
            for index, run in enumerate(paragraph.runs):
                next_cursor = cursor + len(run.text)
                if start_run is None and start < next_cursor:
                    start_run = index
                    start_offset = start - cursor
                if end <= next_cursor:
                    end_run = index
                    end_offset = end - cursor
                    break
                cursor = next_cursor
            if start_run is None or end_run is None:
                break
            first = paragraph.runs[start_run]
            last = paragraph.runs[end_run]
            prefix = first.text[:start_offset]
            suffix = last.text[end_offset:]
            if start_run == end_run:
                first.text = prefix + value + suffix
            else:
                first.text = prefix + value
                for index in range(start_run + 1, end_run):
                    paragraph.runs[index].text = ""
                last.text = suffix


def _replace_template_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_template_paragraph(paragraph, replacements)
            for nested in cell.tables:
                _replace_template_table(nested, replacements)


def fill_writing_template(document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_template_paragraph(paragraph, replacements)
    for table in document.tables:
        _replace_template_table(table, replacements)
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in container.paragraphs:
                _replace_template_paragraph(paragraph, replacements)
            for table in container.tables:
                _replace_template_table(table, replacements)


def create_writing(
    case: Case,
    title: str,
    template: Path | None = None,
    professional: str = "",
) -> Path:
    stem = f"{date.today().isoformat()} - {safe_name(title).upper()}"
    path = unique_path(case.path / f"{stem}.docx")
    if template:
        if not template.is_file():
            raise FileNotFoundError("No encontramos el modelo seleccionado.")
        shutil.copy2(template, path)
        from docx import Document

        document = Document(path)
        _set_spanish_argentina(document)
        fill_writing_template(document, writing_template_values(case, title, professional))
        document.save(path)
        return path

    from docx import Document

    document = Document()
    _set_spanish_argentina(document)
    section = document.sections[0]
    section.top_margin = section.bottom_margin = 1134000
    paragraph = document.add_paragraph()
    paragraph.alignment = 2
    run = paragraph.add_run("{{TITULO}}.")
    run.bold = True
    caption = document.add_paragraph()
    caption.alignment = 2
    caption_run = caption.add_run("{{CARATULA}}")
    caption_run.bold = True
    document.add_paragraph("Señor/a Juez/a:")
    document.add_paragraph("[Complete aquí el escrito]")
    fill_writing_template(document, writing_template_values(case, title, professional))
    document.save(path)
    return path


def list_models(models_dir: Path) -> list[Path]:
    if not models_dir.is_dir():
        return []
    return sorted(
        (path for path in models_dir.iterdir() if path.suffix.lower() == ".docx"),
        key=lambda path: path.name.casefold(),
    )


def add_model(models_dir: Path, source: Path) -> Path:
    models_dir.mkdir(parents=True, exist_ok=True)
    target = unique_path(models_dir / normalize_filename(source.name, ".docx"))
    shutil.copy2(source, target)
    return target


def _office_cache_paths(source: Path) -> tuple[Path, Path]:
    key = hashlib.sha256(str(source.resolve()).casefold().encode("utf-8")).hexdigest()
    return OFFICE_CACHE_DIR / f"{key}.pdf", OFFICE_CACHE_DIR / f"{key}.json"


def _office_signature(source: Path) -> dict[str, int]:
    stat = source.stat()
    return {"size": stat.st_size, "modified": stat.st_mtime_ns}


def _restore_office_cache(source: Path, output: Path) -> bool:
    cached_pdf, cached_metadata = _office_cache_paths(source)
    try:
        metadata = json.loads(cached_metadata.read_text(encoding="utf-8"))
        if metadata != _office_signature(source) or not cached_pdf.is_file():
            return False
        shutil.copy2(cached_pdf, output)
        return output.stat().st_size > 0
    except (OSError, ValueError, TypeError):
        return False


def _store_office_cache(source: Path, output: Path) -> None:
    try:
        OFFICE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        cached_pdf, cached_metadata = _office_cache_paths(source)
        shutil.copy2(output, cached_pdf)
        cached_metadata.write_text(
            json.dumps(_office_signature(source)),
            encoding="utf-8",
        )
    except OSError:
        # The cache only improves speed; a permission issue must not block compilation.
        pass


def _terminate_process_tree(process: subprocess.Popen) -> None:
    if process.poll() is not None:
        return
    if os.name == "nt":
        subprocess.run(
            ["taskkill.exe", "/PID", str(process.pid), "/T", "/F"],
            check=False,
            capture_output=True,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
    else:
        process.terminate()
    try:
        process.wait(timeout=5)
    except subprocess.TimeoutExpired:
        process.kill()


def _run_cancellable_process(
    command: list[str],
    timeout: float,
    cancelled: Callable[[], bool],
) -> None:
    process = subprocess.Popen(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    deadline = time.monotonic() + timeout
    try:
        while process.poll() is None:
            _check_cancelled(cancelled)
            if time.monotonic() >= deadline:
                raise subprocess.TimeoutExpired(command, timeout)
            time.sleep(0.1)
        stdout, stderr = process.communicate()
        if process.returncode:
            raise subprocess.CalledProcessError(process.returncode, command, stdout, stderr)
    except (CompilationCancelled, subprocess.TimeoutExpired):
        _terminate_process_tree(process)
        raise


def _office_to_pdf(
    source: Path,
    destination: Path,
    cancelled: Callable[[], bool] = lambda: False,
) -> Path:
    """Convert outside the app process so a Word/COM failure cannot freeze the UI."""
    output = destination / f"{source.stem}.pdf"
    _check_cancelled(cancelled)
    if _restore_office_cache(source, output):
        return output

    if os.name == "nt":
        def quote(path: Path) -> str:
            return "'" + str(path.resolve()).replace("'", "''") + "'"

        script = (
            "$ErrorActionPreference='Stop'; $word=$null; $doc=$null; "
            "try { $word=New-Object -ComObject Word.Application; $word.Visible=$false; "
            f"$doc=$word.Documents.Open({quote(source)}, $false, $true); "
            f"$doc.ExportAsFixedFormat({quote(output)}, 17); "
            "} finally { if($doc -ne $null){$doc.Close($false)}; "
            "if($word -ne $null){$word.Quit()} }"
        )
        try:
            _run_cancellable_process(
                ["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", script],
                45,
                cancelled,
            )
            if output.exists():
                _store_office_cache(source, output)
                return output
        except (subprocess.SubprocessError, OSError):
            pass

    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        Path(os.getenv("PROGRAMFILES", "")) / "LibreOffice" / "program" / "soffice.exe",
        Path(os.getenv("PROGRAMFILES(X86)", "")) / "LibreOffice" / "program" / "soffice.exe",
    ]
    executable = next(
        (str(candidate) for candidate in candidates if candidate and Path(candidate).is_file()),
        None,
    )
    if not executable:
        raise RuntimeError(
            "No pudimos usar Microsoft Word para convertir el archivo. "
            "Comprobá que Word abra normalmente o instalá LibreOffice."
        )
    try:
        _run_cancellable_process(
            [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(destination),
                str(source),
            ],
            120,
            cancelled,
        )
    except subprocess.SubprocessError as exc:
        raise RuntimeError("LibreOffice no pudo convertir el archivo.") from exc
    if not output.exists():
        raise RuntimeError("La conversión terminó sin generar el PDF esperado.")
    _store_office_cache(source, output)
    return output


def to_pdf(
    source: Path,
    destination: Path,
    cancelled: Callable[[], bool] = lambda: False,
) -> Path:
    destination.mkdir(parents=True, exist_ok=True)
    _check_cancelled(cancelled)
    extension = source.suffix.lower()
    if extension == ".pdf":
        output = destination / source.name
        import pymupdf

        document = pymupdf.open(source)
        try:
            encryption = str(document.metadata.get("encryption", "")).strip()
            if document.needs_pass and document.authenticate("") <= 0:
                raise RuntimeError(
                    f"{source.name} está protegido con contraseña. "
                    "Abrilo y guardá una copia sin contraseña antes de compilar."
                )
            if encryption:
                # SRT and other portals may emit readable PDFs with AES encryption.
                # pypdf needs optional crypto support for those streams, so create
                # an unencrypted temporary copy without changing the case file.
                document.save(
                    output,
                    encryption=pymupdf.PDF_ENCRYPT_NONE,
                    garbage=4,
                    deflate=True,
                )
            elif source.resolve() != output.resolve():
                shutil.copy2(source, output)
        finally:
            document.close()
        return output
    if extension in IMAGE_EXTENSIONS:
        from PIL import Image, ImageOps

        output = destination / f"{source.stem}.pdf"
        with Image.open(source) as image:
            frames = []
            for index in range(getattr(image, "n_frames", 1)):
                image.seek(index)
                frames.append(ImageOps.exif_transpose(image.copy()).convert("RGB"))
            frames[0].save(
                output,
                "PDF",
                save_all=True,
                append_images=frames[1:],
                resolution=150,
            )
        return output
    if extension in OFFICE_EXTENSIONS:
        return _office_to_pdf(source, destination, cancelled)
    raise RuntimeError(f"No podemos convertir el formato {source.suffix or '(sin extensión)' }.")


def merge_pdfs(
    paths: list[Path],
    output: Path,
    cancelled: Callable[[], bool] = lambda: False,
):
    from io import BytesIO

    from pypdf import PdfReader, PdfWriter
    from pypdf.errors import DependencyError

    writer = PdfWriter()
    # PdfReader(path) keeps a Windows file handle alive through lazy PDF
    # objects. Work from in-memory streams so every case file is released as
    # soon as read, even while PdfWriter is still composing the result.
    sources: list[BytesIO] = []
    try:
        try:
            for path in paths:
                _check_cancelled(cancelled)
                source = BytesIO(path.read_bytes())
                sources.append(source)
                for page in PdfReader(source).pages:
                    _check_cancelled(cancelled)
                    writer.add_page(page)
        except DependencyError as error:
            raise RuntimeError(
                "Uno de los PDFs usa cifrado AES. Actualizá el soporte de PDF "
                "o guardá una copia sin protección antes de compilar."
            ) from error
        with output.open("wb") as stream:
            writer.write(stream)
    finally:
        if hasattr(writer, "close"):
            writer.close()
        for source in sources:
            source.close()


def compress_pdf(
    source: Path,
    output: Path,
    limit: int | None = None,
    progress: Callable[[str], None] = lambda _: None,
    cancelled: Callable[[], bool] = lambda: False,
):
    """Optimize embedded images while preserving selectable text and page geometry."""
    import pymupdf

    source_size = max(1, source.stat().st_size)
    _check_cancelled(cancelled)
    ratio = min(1.0, (limit or source_size) / source_size)
    if ratio >= 0.75:
        target_dpi, quality = 130, 65
    elif ratio >= 0.55:
        target_dpi, quality = 110, 55
    else:
        target_dpi, quality = 90, 45

    fast_output = output.with_name(f"{output.stem}-rapido{output.suffix}")
    document = pymupdf.open(source)
    try:
        progress("Optimizando la estructura del PDF…")
        document.save(
            fast_output,
            garbage=4,
            clean=True,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
            use_objstms=1,
            compression_effort=100,
        )
    finally:
        document.close()

    if limit and fast_output.stat().st_size <= limit:
        _check_cancelled(cancelled)
        os.replace(fast_output, output)
        return

    _check_cancelled(cancelled)
    candidate = output.with_name(f"{output.stem}-imagenes{output.suffix}")
    document = pymupdf.open(source)
    try:
        progress("Optimizando las imágenes sin convertir las páginas en fotos…")
        if hasattr(document, "rewrite_images"):
            document.rewrite_images(
                dpi_threshold=target_dpi + 25,
                dpi_target=target_dpi,
                quality=quality,
            )
        document.save(
            candidate,
            garbage=4,
            clean=True,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
            use_objstms=1,
            compression_effort=100,
        )
    finally:
        document.close()

    selected = candidate if candidate.stat().st_size < fast_output.stat().st_size else fast_output
    _check_cancelled(cancelled)
    os.replace(selected, output)
    for temporary in (candidate, fast_output):
        if temporary.exists():
            temporary.unlink()


def compile_documents(
    case: Case,
    paths: list[Path],
    limit: int,
    output_name: str,
    progress: Callable[[str], None] = lambda _: None,
    cancelled: Callable[[], bool] = lambda: False,
    replace_existing: bool = False,
) -> CompilationResult:
    _check_cancelled(cancelled)
    if not paths:
        raise ValueError("Agregá al menos un archivo para compilar.")
    for path in paths:
        if not path.is_file():
            raise FileNotFoundError(f"No encontramos {path.name}.")
        if path.suffix.lower() not in PDF_EXTENSIONS:
            raise ValueError(f"{path.name} no se puede incluir en un PDF.")

    case.ensure()
    with tempfile.TemporaryDirectory(prefix="gestor-documental-") as directory:
        temporary = Path(directory)
        converted = []
        for index, path in enumerate(paths, 1):
            _check_cancelled(cancelled)
            progress(f"Preparando {index} de {len(paths)} · {path.name}")
            converted.append(to_pdf(path, temporary / f"item-{index:03d}", cancelled))

        _check_cancelled(cancelled)
        merged = temporary / "compilado.pdf"
        progress("Uniendo los archivos en el orden elegido…")
        merge_pdfs(converted, merged, cancelled)
        selected = merged
        compressed = False
        if merged.stat().st_size > limit:
            progress("Reduciendo el tamaño del PDF…")
            compressed_path = temporary / "compilado-comprimido.pdf"
            compress_pdf(merged, compressed_path, limit, progress, cancelled)
            if compressed_path.stat().st_size < merged.stat().st_size:
                selected = compressed_path
                compressed = True

        target_name = normalize_filename(output_name, ".pdf")
        requested_target = case.path / target_name
        target = requested_target if replace_existing else versioned_path(requested_target)
        _check_cancelled(cancelled)
        if replace_existing and target.exists():
            staged = case.path / f".gestor-nuevo-{time.time_ns()}.pdf"
            try:
                shutil.copy2(selected, staged)
                move_to_recycle_bin([target])
                os.replace(staged, target)
            finally:
                if staged.exists():
                    staged.unlink()
        else:
            shutil.copy2(selected, target)
        return CompilationResult(
            output=target,
            limit=limit,
            compressed=compressed,
            exceeds_limit=target.stat().st_size > limit,
        )


def split_pdf(source: Path, limit: int, output_dir: Path, stem: str | None = None) -> list[Path]:
    from io import BytesIO

    from pypdf import PdfReader, PdfWriter

    # Keep the source in memory so Windows can rename or remove the case
    # folder immediately after the split finishes.
    source_stream = BytesIO(source.read_bytes())
    reader = PdfReader(source_stream)
    results = []
    writer = PdfWriter()
    part = 1
    output_stem = safe_name(stem or source.stem)

    def save(current: PdfWriter, number: int) -> Path:
        path = unique_path(output_dir / f"{output_stem} - PARTE {number}.pdf")
        with path.open("wb") as stream:
            current.write(stream)
        return path

    try:
        for page in reader.pages:
            trial = PdfWriter()
            buffer = BytesIO()
            try:
                for previous in writer.pages:
                    trial.add_page(previous)
                trial.add_page(page)
                trial.write(buffer)
                if len(buffer.getvalue()) > limit and len(writer.pages):
                    results.append(save(writer, part))
                    part += 1
                    if hasattr(writer, "close"):
                        writer.close()
                    writer = PdfWriter()
                writer.add_page(page)
            finally:
                if hasattr(trial, "close"):
                    trial.close()
                buffer.close()
        if len(writer.pages):
            results.append(save(writer, part))
        return results
    finally:
        if hasattr(writer, "close"):
            writer.close()
        source_stream.close()


def focus_or_launch_signer(executable: Path):
    if not executable.is_file():
        raise FileNotFoundError("No encontramos la aplicación de firma configurada.")
    if os.name == "nt":
        import ctypes
        from ctypes import wintypes

        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        process_query_limited_information = 0x1000
        found: list[int] = []

        def enum_callback(window_handle, _):
            if not user32.IsWindowVisible(window_handle):
                return True
            process_id = wintypes.DWORD()
            user32.GetWindowThreadProcessId(window_handle, ctypes.byref(process_id))
            process = kernel32.OpenProcess(
                process_query_limited_information,
                False,
                process_id.value,
            )
            if not process:
                return True
            try:
                size = wintypes.DWORD(32768)
                buffer = ctypes.create_unicode_buffer(size.value)
                if kernel32.QueryFullProcessImageNameW(process, 0, buffer, ctypes.byref(size)):
                    if Path(buffer.value).resolve() == executable.resolve():
                        found.append(window_handle)
                        return False
            finally:
                kernel32.CloseHandle(process)
            return True

        callback = ctypes.WINFUNCTYPE(wintypes.BOOL, wintypes.HWND, wintypes.LPARAM)(enum_callback)
        user32.EnumWindows(callback, 0)
        if found:
            user32.ShowWindow(found[0], 9)
            user32.SetForegroundWindow(found[0])
            return
    subprocess.Popen([str(executable)])



def human_size(value: int) -> str:
    if value < 1024 * 1024:
        return f"{value / 1024:.0f} KB"
    return f"{value / 1024 / 1024:.2f} MB"
