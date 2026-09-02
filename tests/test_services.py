import tempfile
import unittest
import json
import zipfile
from datetime import date
from pathlib import Path
from unittest.mock import patch

from pypdf import PdfReader, PdfWriter
import pymupdf
import gestor_documental.services as services

from gestor_documental.models import (
    BUENOS_AIRES_LIMIT,
    SISFE_COMMON_LIMIT,
    SISFE_SPECIAL_LIMIT,
    SRT_LIMIT,
    limit_for,
)
from gestor_documental.services import (
    CompilationCancelled,
    SettingsStore,
    case_matches,
    compile_documents,
    compress_pdf,
    create_case,
    create_writing,
    ensure_default_writing_template,
    ensure_bundled_writing_models,
    import_file,
    import_directory,
    list_cases,
    normalize_filename,
    read_case_metadata,
    repair_text,
    rename_case,
    rename_case_entry,
    rename_case_file,
    safe_name,
    save_case_metadata,
    short_case_identifier,
    suggested_presentation_name,
    spanish_long_date,
    split_pdf,
    study_library_path,
    versioned_path,
    writing_template_values,
)


def make_pdf(path: Path, pages: int = 1):
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=595, height=842)
    with path.open("wb") as stream:
        writer.write(stream)


class ServiceTests(unittest.TestCase):
    def test_case_starts_empty_without_automatic_subfolders(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory) / "Estudio"
            study.mkdir()
            case = create_case(study, "Gómez c/ SIJAM")
            self.assertTrue(case.path.is_dir())
            self.assertEqual(list(case.path.iterdir()), [])
            self.assertEqual(list_cases(study), [case])

    def test_user_folders_are_recognized_imported_and_renamed(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            case = create_case(root / "Estudio", "Caso")
            existing = case.path / "Documental del actor"
            existing.mkdir()
            make_pdf(existing / "DNI.pdf")
            self.assertIn(existing, case.entries())
            self.assertIn(existing / "DNI.pdf", case.files())

            source = root / "Carpeta importada"
            nested = source / "Nivel 2"
            nested.mkdir(parents=True)
            make_pdf(nested / "Prueba.pdf")
            imported = import_directory(case, source)
            self.assertTrue((imported / "Nivel 2" / "Prueba.pdf").is_file())

            renamed = rename_case_entry(imported, "Informes")
            self.assertEqual(renamed.name, "Informes")
            self.assertTrue((renamed / "Nivel 2" / "Prueba.pdf").is_file())

    def test_quick_access_folder_is_not_treated_as_a_case(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory) / "Estudio"
            study.mkdir()
            case = create_case(study, "Caso")
            library = study_library_path(study, create=True)
            self.assertTrue(library.is_dir())
            self.assertEqual(list_cases(study), [case])

    def test_case_metadata_and_smart_search(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Rosales c/ Provincia")
            save_case_metadata(
                case,
                {
                    "Actor": "Pablo Rosales",
                    "CUIJ": "21-12345678-9",
                    "Radicación": "Juzgado Laboral Nº 4",
                },
            )
            save_case_metadata(
                case,
                {
                    "Actor": "Pablo Rosales",
                    "Causa": "DAÑOS Y PERJUICIOS",
                    "CUIJ": "21-12345678-9",
                    "Radicación": "Juzgado Laboral Nº 4",
                },
            )
            self.assertEqual(read_case_metadata(case)["Actor"], "Pablo Rosales")
            self.assertEqual(read_case_metadata(case)["Causa"], "DAÑOS Y PERJUICIOS")
            self.assertTrue(case_matches(case, "rosales 12345678"))
            self.assertTrue(case_matches(case, "laboral numero 4"))
            self.assertFalse(case_matches(case, "gonzalez"))

    def test_presentation_name_uses_actor_date_and_writing_title(self):
        with tempfile.TemporaryDirectory() as directory:
            case = create_case(Path(directory), "Caso")
            save_case_metadata(case, {"Actor": "YOCCA, EMANUEL Y OT."})
            writing = case.path / "2026-08-13 - DEMANDA.docx"
            self.assertEqual(short_case_identifier(case), "YOCCA")
            self.assertEqual(
                suggested_presentation_name(case, writing, date(2026, 8, 21)),
                "YOCCA_2026-08-21_DEMANDA.pdf",
            )

    def test_short_name_and_custom_metadata_are_available_to_models(self):
        with tempfile.TemporaryDirectory() as directory:
            case = create_case(Path(directory), "Caso")
            save_case_metadata(
                case,
                {
                    "Actor": "Pérez, Ana",
                    "Nombre corto para archivos": "PEREZ-ART",
                    "Jurisdicción": "Santa Fe",
                    "Nombre del mediador": "María López",
                },
            )
            values = writing_template_values(case, "Apela")
            self.assertEqual(values["{{NOMBRE_CORTO}}"], "PEREZ-ART")
            self.assertEqual(values["{{JURISDICCION}}"], "Santa Fe")
            self.assertEqual(values["{{NOMBRE_DEL_MEDIADOR}}"], "María López")

    def test_new_case_data_and_calculations_are_available_to_word_models(self):
        with tempfile.TemporaryDirectory() as directory:
            case = create_case(Path(directory), "Caso")
            save_case_metadata(
                case,
                {
                    "Actor": "Pérez, Ana",
                    "DNI del actor": "30111222",
                    "Domicilio real": "Calle 1",
                    "Fecha de nacimiento": "10/05/1990",
                    "Fecha del accidente": "10/05/2026",
                    "Fecha de ingreso": "10/05/2020",
                    "Posibles testigos": "María | 341 111\nLuis | 341 222",
                },
            )
            values = writing_template_values(case, "Demanda")
            self.assertEqual(values["{{DOCUMENTO_ACTOR}}"], "30111222")
            self.assertEqual(values["{{DOMICILIO_ACTOR}}"], "Calle 1")
            self.assertEqual(values["{{EDAD_RAEO}}"], "36")
            self.assertEqual(values["{{ANTIGUEDAD_LABORAL}}"], "6 años")
            self.assertEqual(
                values["{{POSIBLES_TESTIGOS}}"],
                "María | 341 111\nLuis | 341 222",
            )

    def test_settings_roundtrip(self):
        with tempfile.TemporaryDirectory() as directory:
            app_dir = Path(directory) / "app"
            study = Path(directory) / "Estudio"
            store = SettingsStore(app_dir)
            self.assertEqual(store.settings.current_professional, "Profesional")
            store.set_study_root(study)
            store.add_professional("Dra. Ana Pérez")
            store.set_layout_state(
                {
                    "body": [220, 1180],
                    "workspace": [900, 300],
                    "compilation_visible": False,
                }
            )
            reloaded = SettingsStore(app_dir)
            self.assertEqual(reloaded.settings.study_root, study)
            self.assertIn("Dra. Ana Pérez", reloaded.settings.professionals)
            self.assertEqual(reloaded.settings.current_professional, "Dra. Ana Pérez")
            self.assertEqual(reloaded.settings.layout_state["body"], [220, 1180])
            self.assertFalse(reloaded.settings.layout_state["compilation_visible"])

    def test_settings_support_multiple_study_locations_and_safe_removal(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            local = root / "Estudio local"
            shared = root / "Google Drive compartido"
            store = SettingsStore(root / "app")
            store.add_study_root(local)
            store.add_study_root(shared)

            reloaded = SettingsStore(root / "app")
            self.assertEqual(reloaded.settings.study_roots, [local, shared])
            self.assertEqual(reloaded.settings.study_root, shared)

            reloaded.remove_study_root(shared)
            self.assertEqual(reloaded.settings.study_roots, [local])
            self.assertEqual(reloaded.settings.study_root, local)
            self.assertTrue(local.is_dir())
            self.assertTrue(shared.is_dir())

    def test_old_single_study_setting_migrates_to_locations(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            app_dir = root / "app"
            app_dir.mkdir()
            study = root / "Estudio anterior"
            (app_dir / "config.json").write_text(
                json.dumps({"study_root": str(study)}),
                encoding="utf-8",
            )
            store = SettingsStore(app_dir)
            self.assertEqual(store.settings.study_roots, [study])
            self.assertEqual(store.settings.study_root, study)

    def test_cedula_labvc_is_installed_once_without_overwriting_user_changes(self):
        with tempfile.TemporaryDirectory() as directory:
            app_dir = Path(directory) / "app"
            store = SettingsStore(app_dir)
            model = store.models_dir / "Cedula LABVC.docx"
            self.assertTrue(model.is_file())
            model.write_bytes(b"modelo personalizado")
            ensure_bundled_writing_models(store.models_dir)
            self.assertEqual(model.read_bytes(), b"modelo personalizado")

    def test_settings_infers_study_from_previous_case_list(self):
        with tempfile.TemporaryDirectory() as directory:
            app_dir = Path(directory) / "app"
            app_dir.mkdir()
            study = Path(directory) / "Estudio"
            legacy_case = study / "Caso anterior"
            (app_dir / "config.json").write_text(
                json.dumps({"cases": [{"name": "Caso anterior", "path": str(legacy_case)}]}),
                encoding="utf-8",
            )
            store = SettingsStore(app_dir)
            self.assertEqual(store.settings.study_root, study)

    def test_settings_does_not_treat_application_directory_as_study(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            app_dir = root / "appdata"
            app_dir.mkdir()
            source_root = root / "Gestor"
            (source_root / "gestor_documental").mkdir(parents=True)
            (source_root / "run.py").write_text("", encoding="utf-8")
            (app_dir / "config.json").write_text(
                json.dumps({"cases": [{"name": "Caso", "path": str(source_root / "Caso")}]}),
                encoding="utf-8",
            )
            store = SettingsStore(app_dir)
            self.assertIsNone(store.settings.study_root)

    def test_import_and_rename_stay_in_case_root(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            study = root / "Estudio"
            study.mkdir()
            case = create_case(study, "Caso")
            source = root / "DNI actor.PDF"
            make_pdf(source)
            imported = import_file(case, source, "DNI: actor.PDF")
            self.assertEqual(imported.parent, case.path)
            self.assertEqual(imported.name, "DNI- actor.pdf")
            renamed = rename_case_file(imported, "01 - DNI ACTOR")
            self.assertEqual(renamed.name, "01 - DNI ACTOR.pdf")

    def test_case_folder_can_be_renamed(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Caso anterior")
            renamed = rename_case(case, "Caso nuevo")
            self.assertFalse(case.path.exists())
            self.assertEqual(renamed.path, study / "Caso nuevo")
            self.assertTrue(renamed.path.is_dir())

    def test_image_import_converts_to_pdf_when_requested(self):
        with tempfile.TemporaryDirectory() as directory:
            from PIL import Image

            root = Path(directory)
            case = create_case(root, "Caso")
            source = root / "foto.png"
            Image.new("RGB", (50, 50), "white").save(source)
            imported = import_file(case, source, "Documento", convert_to_pdf=True)
            self.assertEqual(imported.suffix, ".pdf")
            self.assertEqual(len(PdfReader(str(imported)).pages), 1)

    def test_image_import_supports_grayscale_and_black_white(self):
        with tempfile.TemporaryDirectory() as directory:
            from PIL import Image
            import pymupdf

            root = Path(directory)
            case = create_case(root, "Caso")
            source = root / "foto.png"
            image = Image.new("RGB", (40, 20), "red")
            image.paste("white", (20, 0, 40, 20))
            image.save(source)

            grayscale = import_file(
                case, source, "Grises", convert_to_pdf=True, image_mode="grayscale"
            )
            black_white = import_file(
                case, source, "Blanco y negro", convert_to_pdf=True, image_mode="black_white"
            )

            minima = []
            for pdf in (grayscale, black_white):
                document = pymupdf.open(pdf)
                try:
                    image_info = document[0].get_images(full=True)[0]
                    extracted = document.extract_image(image_info[0])["image"]
                finally:
                    document.close()
                from io import BytesIO
                with Image.open(BytesIO(extracted)) as converted:
                    self.assertEqual(converted.mode, "L")
                    minima.append(converted.getextrema()[0])
            self.assertGreater(minima[0], 0)
            self.assertEqual(minima[1], 0)

    def test_writing_is_created_directly_in_case(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Caso")
            writing = create_writing(case, "Demanda")
            self.assertEqual(writing.parent, case.path)
            self.assertTrue(writing.is_file())
            self.assertFalse(any(path.is_dir() for path in case.path.iterdir()))

    def test_editable_base_template_is_copied_and_uses_spanish_argentina(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            template = ensure_default_writing_template(root / "Modelos" / "Modelo base.docx")
            with zipfile.ZipFile(template) as package:
                styles = package.read("word/styles.xml").decode("utf-8")
                settings = package.read("word/settings.xml").decode("utf-8")
            self.assertIn('w:val="es-AR"', styles)
            self.assertIn('w:val="es-AR"', settings)

            from docx import Document

            document = Document(template)
            document.add_paragraph("CAMBIO EXTERNO")
            document.save(template)
            case = create_case(root, "Caso")
            writing = create_writing(case, "Escrito", template)
            copied = Document(writing)
            self.assertIn("CAMBIO EXTERNO", [paragraph.text for paragraph in copied.paragraphs])

    def test_previous_generated_base_is_upgraded_to_supplied_office_template(self):
        with tempfile.TemporaryDirectory() as directory:
            from docx import Document

            path = Path(directory) / "Modelo base - Escrito nuevo.docx"
            previous = Document()
            previous.add_paragraph("TÍTULO DEL ESCRITO.")
            previous.add_paragraph("Señor/a Juez/a:")
            previous.add_paragraph("[Complete aquí el escrito]")
            previous.save(path)

            upgraded = ensure_default_writing_template(path)
            paragraphs = [paragraph.text for paragraph in Document(upgraded).paragraphs]
            self.assertEqual(len(paragraphs), 11)
            self.assertIn("{{PROFESIONAL}}", paragraphs[3])
            self.assertIn("{{CARATULA}}", paragraphs[3])

    def test_base_writing_uses_selected_professional_caption_and_case_metadata(self):
        with tempfile.TemporaryDirectory() as directory:
            from docx import Document

            root = Path(directory)
            case = create_case(root / "Estudio", "Caso")
            save_case_metadata(
                case,
                {
                    "Actor": "Pérez, Ana",
                    "Demandado": "Acme S.A.",
                    "Causa": "Daños y perjuicios",
                    "CUIJ": "21-12345678-9",
                },
            )
            template = ensure_default_writing_template(root / "Modelos" / "Modelo base.docx")
            writing = create_writing(case, "Manifiesta", template, "Dra. Laura Gómez")
            paragraphs = [paragraph.text for paragraph in Document(writing).paragraphs]
            presentation = paragraphs[3]
            self.assertIn("LAURA GÓMEZ, abogado", presentation)
            self.assertIn(
                '“PÉREZ, ANA C/ ACME S.A. S/ DAÑOS Y PERJUICIOS” '
                "(CUIJ N° 21-12345678-9)",
                presentation,
            )
            self.assertNotIn("{{", "\n".join(paragraphs))

    def test_writing_prefers_abogado_metadata_in_uppercase_without_title(self):
        with tempfile.TemporaryDirectory() as directory:
            from docx import Document

            root = Path(directory)
            case = create_case(root / "Estudio", "Caso")
            save_case_metadata(case, {"Abogado": "Dr. Julián Nicolás Costa"})
            template = ensure_default_writing_template(root / "Modelo.docx")
            writing = create_writing(case, "Apela", template, "Dra. Otra Persona")
            presentation = Document(writing).paragraphs[3].text
            self.assertTrue(presentation.startswith("JULIÁN NICOLÁS COSTA, abogado"))
            self.assertNotIn("DR.", presentation.upper())

    def test_cedula_labvc_fills_long_date_and_case_metadata(self):
        with tempfile.TemporaryDirectory() as directory:
            from docx import Document

            root = Path(directory)
            case = create_case(root / "Estudio", "Caso")
            save_case_metadata(
                case,
                {
                    "Actor": "Pérez, Ana",
                    "Demandado": "Empresa S.A.",
                    "Causa": "Cobro de pesos",
                    "CUIJ": "21-12345678-9",
                },
            )
            model = (
                Path(services.__file__).with_name("bundled-models")
                / "Cedula LABVC.docx"
            )
            writing = create_writing(case, "Cedula LABVC", model)
            document = Document(writing)
            self.assertEqual(
                document.paragraphs[5].text,
                f"Villa Constitución, {spanish_long_date(date.today())}",
            )
            self.assertEqual(document.paragraphs[13].text, "Por\t\tPérez, Ana")
            self.assertEqual(document.paragraphs[14].text, "Contra\tEmpresa S.A.")
            self.assertEqual(document.paragraphs[15].text, "Sobre\t\tCobro de pesos")
            self.assertEqual(
                document.paragraphs[16].text,
                "Número de expediente\t21-12345678-9",
            )
            self.assertEqual(len(document.tables), 1)
            with zipfile.ZipFile(writing) as package:
                xml = package.read("word/document.xml").decode("utf-8")
                self.assertEqual(xml.count("FORMTEXT"), 5)
                self.assertIn("word/media/image1.jpeg", package.namelist())
            self.assertEqual(spanish_long_date(date(2026, 8, 13)), "13 de agosto de 2026")

    def test_compile_can_be_cancelled_without_creating_output(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            case = create_case(root, "Caso")
            source = case.path / "documental.pdf"
            make_pdf(source)
            with self.assertRaises(CompilationCancelled):
                compile_documents(
                    case,
                    [source],
                    BUENOS_AIRES_LIMIT,
                    "No debe existir",
                    cancelled=lambda: True,
                )
            self.assertFalse((case.path / "No debe existir.pdf").exists())

    def test_office_conversion_cache_invalidates_when_source_changes(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "escrito.docx"
            source.write_bytes(b"version uno")
            converted = root / "convertido.pdf"
            make_pdf(converted)
            with patch.object(services, "OFFICE_CACHE_DIR", root / "cache"):
                services._store_office_cache(source, converted)
                restored = root / "restaurado.pdf"
                self.assertTrue(services._restore_office_cache(source, restored))
                self.assertEqual(restored.read_bytes(), converted.read_bytes())
                source.write_bytes(b"version dos con cambios")
                self.assertFalse(services._restore_office_cache(source, root / "obsoleto.pdf"))

    def test_pdf_image_optimization_preserves_pages(self):
        with tempfile.TemporaryDirectory() as directory:
            from PIL import Image

            root = Path(directory)
            source = root / "imagenes.pdf"
            output = root / "optimizado.pdf"
            Image.new("RGB", (1600, 2200), "white").save(source, "PDF", quality=95)
            compress_pdf(source, output, 100_000)
            self.assertTrue(output.is_file())
            self.assertEqual(len(PdfReader(str(output)).pages), 1)

    def test_compile_creates_one_ordered_pdf_in_case(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Caso")
            first = case.path / "01 documental.pdf"
            writing = case.path / "02 escrito.pdf"
            make_pdf(first, 1)
            make_pdf(writing, 2)
            result = compile_documents(
                case,
                [first, writing],
                BUENOS_AIRES_LIMIT,
                "Presentación prueba",
            )
            self.assertEqual(result.output.parent, case.path)
            self.assertEqual(len(PdfReader(str(result.output)).pages), 3)
            self.assertFalse(result.exceeds_limit)

    def test_compilation_versions_are_clear_and_existing_pdf_can_be_replaced(self):
        with tempfile.TemporaryDirectory() as directory:
            case = create_case(Path(directory), "Caso")
            source = case.path / "documental.pdf"
            make_pdf(source)
            first = compile_documents(
                case,
                [source],
                BUENOS_AIRES_LIMIT,
                "PEREZ_2026-08-21_DEMANDA.pdf",
            )
            self.assertEqual(
                versioned_path(first.output).name,
                "PEREZ_2026-08-21_DEMANDA_V2.pdf",
            )

            def recycle(paths):
                for path in paths:
                    path.unlink()

            with patch("gestor_documental.services.move_to_recycle_bin", side_effect=recycle):
                replaced = compile_documents(
                    case,
                    [source],
                    BUENOS_AIRES_LIMIT,
                    first.output.name,
                    replace_existing=True,
                )
            self.assertEqual(replaced.output, first.output)
            self.assertTrue(replaced.output.is_file())
            self.assertFalse((case.path / "PEREZ_2026-08-21_DEMANDA_V2.pdf").exists())

    def test_compile_releases_case_files_immediately(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Caso liberable")
            first = case.path / "documental.pdf"
            second = case.path / "escrito.pdf"
            make_pdf(first)
            make_pdf(second)

            compile_documents(
                case,
                [first, second],
                BUENOS_AIRES_LIMIT,
                "Presentación sin bloqueos",
            )

            # Windows refuses this rename if any input or output is still open.
            renamed = study / "Caso liberado"
            case.path.rename(renamed)
            self.assertTrue(renamed.is_dir())

    def test_compile_normalizes_aes_pdf_without_touching_original(self):
        with tempfile.TemporaryDirectory() as directory:
            study = Path(directory)
            case = create_case(study, "Caso AES")
            encrypted = case.path / "archivo SRT.pdf"
            document = pymupdf.open()
            document.new_page()
            document.save(
                encrypted,
                encryption=pymupdf.PDF_ENCRYPT_AES_256,
                owner_pw="propietario",
                user_pw="",
            )
            document.close()
            original = encrypted.read_bytes()

            result = compile_documents(
                case,
                [encrypted],
                BUENOS_AIRES_LIMIT,
                "Compilado AES",
            )

            self.assertTrue(result.output.is_file())
            self.assertEqual(encrypted.read_bytes(), original)
            output_document = pymupdf.open(result.output)
            try:
                self.assertFalse(output_document.metadata.get("encryption"))
                self.assertEqual(output_document.page_count, 1)
            finally:
                output_document.close()

    def test_split_pdf(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "compilado.pdf"
            make_pdf(source, 3)
            parts = split_pdf(source, 600, root)
            self.assertGreaterEqual(len(parts), 1)
            self.assertTrue(all(path.parent == root for path in parts))

    def test_names_and_limits(self):
        self.assertEqual(safe_name('Gómez: c/ "SIJAM"'), "Gómez- c- -SIJAM")
        self.assertEqual(normalize_filename("  Poder   ACTOR.DOCX"), "Poder ACTOR.docx")
        self.assertEqual(limit_for("SRT · 1 MB"), SRT_LIMIT)
        self.assertEqual(limit_for("SISFE común · 3 MB"), SISFE_COMMON_LIMIT)
        self.assertEqual(limit_for("SISFE demanda/contestación · 6 MB"), SISFE_SPECIAL_LIMIT)
        self.assertEqual(limit_for("Provincia de Buenos Aires · 20 MB"), BUENOS_AIRES_LIMIT)
        with self.assertRaises(ValueError):
            limit_for("Sistema inexistente")
        self.assertEqual(repair_text("JuliÃ¡n NicolÃ¡s"), "Julián Nicolás")


if __name__ == "__main__":
    unittest.main()
