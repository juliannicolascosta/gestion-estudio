from __future__ import annotations
import json, os, shutil, subprocess
from datetime import date
from pathlib import Path
from typing import Callable

from .models import Case, COMMON_LIMIT, SPECIAL_LIMIT, SPECIAL_TYPES

APP_DIR = Path(os.getenv("APPDATA", Path.home())) / "PreparadorSISFE"
CONFIG = APP_DIR / "config.json"

def safe_name(value: str) -> str:
    bad = '<>:"/\\|?*'
    return " ".join("".join("-" if c in bad else c for c in value).split()).strip(" .-")

class CaseStore:
    def __init__(self):
        APP_DIR.mkdir(parents=True, exist_ok=True)
        self.cases: list[Case] = []
        self.load()

    def load(self):
        try:
            rows = json.loads(CONFIG.read_text(encoding="utf-8")).get("cases", [])
            self.cases = [Case(r["name"], Path(r["path"])) for r in rows]
        except (FileNotFoundError, ValueError, KeyError, TypeError):
            self.cases = []

    def save(self):
        CONFIG.write_text(json.dumps({"cases": [{"name": c.name, "path": str(c.path)} for c in self.cases]}, ensure_ascii=False, indent=2), encoding="utf-8")

    def add(self, name: str, root: Path) -> Case:
        case = Case(safe_name(name), root / safe_name(name))
        case.ensure()
        self.cases = [c for c in self.cases if c.path != case.path] + [case]
        self.save()
        return case

    def attach(self, path: Path) -> Case:
        case = Case(path.name, path)
        case.ensure()
        self.cases = [c for c in self.cases if c.path != path] + [case]
        self.save()
        return case

def create_writing(case: Case, kind: str) -> Path:
    from docx import Document
    stem = f"{date.today().isoformat()} - {safe_name(kind).upper()}"
    path = case.writings / f"{stem}.docx"
    n = 2
    while path.exists():
        path = case.writings / f"{stem} ({n}).docx"; n += 1
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = 1134000
    p = doc.add_paragraph()
    p.alignment = 2
    r = p.add_run(kind.upper() + ".")
    r.bold = True
    doc.add_paragraph("Señor/a Juez/a:")
    doc.add_paragraph("[Complete aquí el escrito]")
    doc.save(path)
    return path

def open_file(path: Path):
    os.startfile(str(path))

def copy_evidence(case: Case, paths: list[Path]) -> list[Path]:
    copied = []
    for source in paths:
        if not source.is_file(): continue
        target = case.evidence / source.name
        if source.resolve() != target.resolve():
            n = 2
            while target.exists():
                target = case.evidence / f"{source.stem} ({n}){source.suffix}"; n += 1
            shutil.copy2(source, target)
        copied.append(target)
    return copied

def _office_to_pdf(source: Path, dest_dir: Path) -> Path:
    out = dest_dir / f"{source.stem}.pdf"
    # Word automation preserves legal-document layout best on Windows.
    try:
        import win32com.client  # type: ignore
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(source.resolve()), ReadOnly=True)
        doc.ExportAsFixedFormat(str(out.resolve()), 17)
        doc.Close(False); word.Quit()
        return out
    except Exception:
        pass

    # Normal Microsoft Word installations expose COM even when pywin32 is not
    # installed. PowerShell lets the portable build use that interface directly.
    if os.name == "nt":
        def ps_quote(value: Path) -> str:
            return "'" + str(value.resolve()).replace("'", "''") + "'"
        script = (
            "$ErrorActionPreference='Stop'; $word=$null; $doc=$null; "
            "try { $word=New-Object -ComObject Word.Application; $word.Visible=$false; "
            f"$doc=$word.Documents.Open({ps_quote(source)}, $false, $true); "
            f"$doc.ExportAsFixedFormat({ps_quote(out)}, 17); "
            "} finally { if($doc -ne $null){$doc.Close($false)}; "
            "if($word -ne $null){$word.Quit()} }"
        )
        try:
            subprocess.run(
                ["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", script],
                check=True, capture_output=True, text=True, timeout=120,
            )
            if out.exists():
                return out
        except (subprocess.SubprocessError, OSError):
            pass

    candidates = [
        shutil.which("soffice"), shutil.which("libreoffice"),
        Path(os.getenv("PROGRAMFILES", "")) / "LibreOffice" / "program" / "soffice.exe",
        Path(os.getenv("PROGRAMFILES(X86)", "")) / "LibreOffice" / "program" / "soffice.exe",
    ]
    exe = next((str(p) for p in candidates if p and Path(p).is_file()), None)
    if not exe:
        raise RuntimeError(
            "No se pudo usar Microsoft Word para convertir el archivo. "
            "Compruebe que Word abra normalmente o instale LibreOffice."
        )
    try:
        subprocess.run([exe, "--headless", "--convert-to", "pdf", "--outdir", str(dest_dir), str(source)], check=True, capture_output=True, timeout=120)
    except subprocess.SubprocessError as exc:
        raise RuntimeError("LibreOffice no pudo convertir el archivo.") from exc
    if not out.exists():
        raise RuntimeError("La conversión terminó pero no se generó el PDF esperado.")
    return out

def to_pdf(source: Path, dest_dir: Path) -> Path:
    dest_dir.mkdir(parents=True, exist_ok=True)
    ext = source.suffix.lower()
    if ext == ".pdf":
        out = dest_dir / source.name
        if source.resolve() != out.resolve(): shutil.copy2(source, out)
        return out
    if ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        from PIL import Image, ImageOps
        out = dest_dir / f"{source.stem}.pdf"
        with Image.open(source) as im:
            frames = []
            for i in range(getattr(im, "n_frames", 1)):
                im.seek(i); frames.append(ImageOps.exif_transpose(im.copy()).convert("RGB"))
            frames[0].save(out, "PDF", save_all=True, append_images=frames[1:], resolution=150)
        return out
    if ext in {".doc", ".docx", ".odt", ".rtf"}:
        return _office_to_pdf(source, dest_dir)
    raise RuntimeError(f"Formato no compatible: {source.suffix}")

def merge_pdfs(paths: list[Path], output: Path):
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    for path in paths:
        for page in PdfReader(str(path)).pages: writer.add_page(page)
    with output.open("wb") as f: writer.write(f)

def compress_pdf(source: Path, output: Path, quality=65, dpi=120):
    import fitz
    from PIL import Image
    from io import BytesIO
    src = fitz.open(source); dst = fitz.open()
    for page in src:
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        buf = BytesIO(); image.save(buf, "JPEG", quality=quality, optimize=True)
        rect = fitz.paper_rect("a4")
        outpage = dst.new_page(width=rect.width, height=rect.height)
        outpage.insert_image(rect, stream=buf.getvalue(), keep_proportion=True)
    dst.save(output, garbage=4, deflate=True); src.close(); dst.close()

def split_pdf(source: Path, limit: int, output_dir: Path, stem: str) -> list[Path]:
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(str(source)); results=[]; writer=PdfWriter(); part=1
    def save(w, number):
        path = output_dir / f"{stem} - PARTE {number}.pdf"
        with path.open("wb") as f: w.write(f)
        return path
    for page in reader.pages:
        trial = PdfWriter()
        for old in writer.pages: trial.add_page(old)
        trial.add_page(page)
        from io import BytesIO
        b=BytesIO(); trial.write(b)
        if len(b.getvalue()) > limit and len(writer.pages):
            results.append(save(writer, part)); part += 1; writer=PdfWriter()
        writer.add_page(page)
    if len(writer.pages): results.append(save(writer, part))
    return results

def prepare(case: Case, writing: Path, evidence: list[Path], kind: str, progress: Callable[[str], None] = lambda _: None) -> tuple[list[Path], list[str]]:
    case.ensure(); out = case.output
    manifest = out / ".preparador-sisfe.json"
    try:
        previous = json.loads(manifest.read_text(encoding="utf-8")).get("generated", [])
        for name in previous:
            target = out / Path(name).name
            if target.is_file(): target.unlink()
    except (FileNotFoundError, ValueError, TypeError):
        pass
    notes=[]; produced=[]; limit = SPECIAL_LIMIT if kind in SPECIAL_TYPES else COMMON_LIMIT
    progress("Convirtiendo escrito…")
    written_pdf = to_pdf(writing, out)
    final_written = out / f"ESCRITO - {writing.stem}.pdf"
    if written_pdf != final_written: written_pdf.replace(final_written)
    produced.append(final_written)
    converted=[]
    for i, item in enumerate(evidence, 1):
        progress(f"Convirtiendo documental {i} de {len(evidence)}…")
        converted.append(to_pdf(item, out / "_temporales"))
    if converted:
        merged = out / "DOCUMENTAL.pdf"; merge_pdfs(converted, merged)
        produced.append(merged)
    final=[]
    for pdf in produced:
        if pdf.stat().st_size <= limit: final.append(pdf); continue
        progress(f"Comprimiendo {pdf.name}…")
        compressed = out / f"{pdf.stem} - COMPRIMIDO.pdf"
        compress_pdf(pdf, compressed)
        pdf.unlink()
        if compressed.stat().st_size <= limit:
            final.append(compressed); notes.append(f"{compressed.name}: comprimido automáticamente.")
        else:
            parts=split_pdf(compressed, limit, out, pdf.stem)
            compressed.unlink(); final.extend(parts)
            notes.append(f"{pdf.name}: dividido en {len(parts)} partes por superar el límite.")
    shutil.rmtree(out / "_temporales", ignore_errors=True)
    for p in final:
        if p.stat().st_size > limit: notes.append(f"ATENCIÓN: {p.name} aún supera el límite; contiene una página indivisible.")
    manifest.write_text(json.dumps({"generated": [p.name for p in final]}, ensure_ascii=False, indent=2), encoding="utf-8")
    return final, notes

def human_size(value: int) -> str:
    return f"{value / 1024 / 1024:.2f} MB"
